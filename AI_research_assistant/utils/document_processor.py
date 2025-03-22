import os
import uuid
import re
from typing import Dict, Any, Tuple, BinaryIO, Optional
from io import BytesIO

# For PDF handling
import PyPDF2
import pdfplumber

# For DOCX handling
import docx

class DocumentProcessor:
    """
    Simplified document processor that extracts text from research papers
    without complex chunking.
    """
    
    def __init__(self):
        """Initialize the document processor with minimal configuration."""
        self.max_document_size_mb = 15  # Default max size
    
    def process_document(self, uploaded_file) -> Dict[str, Any]:
        """
        Process an uploaded document and extract its full text.
        
        Args:
            uploaded_file: Streamlit uploaded file object
            
        Returns:
            Dictionary with document info and extracted text
        """
        if uploaded_file is None:
            return {"error": "No file uploaded"}
            
        # Check file size
        if uploaded_file.size > self.max_document_size_mb * 1024 * 1024:
            return {"error": f"File too large (max {self.max_document_size_mb}MB)"}
        
        # Generate document ID
        document_id = str(uuid.uuid4())
        
        # Get file extension
        filename = uploaded_file.name
        file_ext = os.path.splitext(filename)[1].lower()
        
        try:
            # Extract text based on file type
            if file_ext == '.pdf':
                text, metadata = self._extract_pdf_text(uploaded_file)
            elif file_ext == '.docx':
                text, metadata = self._extract_docx_text(uploaded_file)
            elif file_ext == '.txt':
                text, metadata = self._extract_text_file(uploaded_file)
            else:
                return {"error": f"Unsupported file format: {file_ext}"}
            
            # Clean the extracted text
            cleaned_text = self._clean_text(text)
            
            # For research papers, try to identify sections
            sections = self._identify_sections(cleaned_text)
            
            return {
                "document_id": document_id,
                "filename": filename,
                "file_type": file_ext,
                "full_text": cleaned_text,
                "metadata": metadata,
                "sections": sections,
                "text_length": len(cleaned_text)
            }
                
        except Exception as e:
            return {"error": f"Failed to process document: {str(e)}"}
    
    def _extract_pdf_text(self, file: BinaryIO) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text from PDF using combined approach for best results.
        
        Args:
            file: File-like object containing PDF data
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        bytes_data = file.getvalue()
        file_stream = BytesIO(bytes_data)
        
        # First try PyPDF2
        text_parts = []
        metadata = {}
        
        try:
            reader = PyPDF2.PdfReader(file_stream)
            
            # Extract metadata
            if reader.metadata:
                for key in reader.metadata:
                    clean_key = key.strip('/').lower()
                    metadata[clean_key] = str(reader.metadata[key])
            
            # Extract text page by page
            for page_num in range(len(reader.pages)):
                page = reader.pages[page_num]
                text_parts.append(page.extract_text() or "")
                
            # If PyPDF2 fails to extract meaningful text, try pdfplumber
            text = "\n".join(text_parts)
            if len(text.strip()) < 100:
                file_stream.seek(0)
                return self._extract_with_pdfplumber(file_stream, metadata)
                
            return text, metadata
                
        except Exception:
            # Fallback to pdfplumber
            file_stream.seek(0)
            return self._extract_with_pdfplumber(file_stream, metadata)
    
    def _extract_with_pdfplumber(self, file_stream: BytesIO, metadata: Dict[str, Any]) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text using pdfplumber (better for complex PDFs).
        
        Args:
            file_stream: BytesIO stream of the PDF
            metadata: Existing metadata dictionary
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        with pdfplumber.open(file_stream) as pdf:
            # Get metadata if not already extracted
            if not metadata and hasattr(pdf, 'metadata') and pdf.metadata:
                metadata = pdf.metadata
                
            # Extract text page by page
            text_parts = []
            for page in pdf.pages:
                text = page.extract_text() or ""
                text_parts.append(text)
                
        return "\n".join(text_parts), metadata
    
    def _extract_docx_text(self, file: BinaryIO) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text from DOCX files.
        
        Args:
            file: File-like object containing DOCX data
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        bytes_data = file.getvalue()
        file_stream = BytesIO(bytes_data)
        
        doc = docx.Document(file_stream)
        
        # Extract all paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():  # Skip empty paragraphs
                paragraphs.append(para.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    paragraphs.append(" | ".join(row_text))
            
        # Extract metadata
        metadata = {}
        if doc.core_properties:
            props = doc.core_properties
            if props.author:
                metadata['author'] = props.author
            if props.title:
                metadata['title'] = props.title
            if props.created:
                metadata['created'] = str(props.created)
            if props.modified:
                metadata['modified'] = str(props.modified)
                
        return "\n".join(paragraphs), metadata
    
    def _extract_text_file(self, file: BinaryIO) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text from plain text files.
        
        Args:
            file: File-like object containing text data
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        text = file.getvalue().decode('utf-8', errors='ignore')
        
        # Basic metadata
        metadata = {
            "filename": file.name
        }
        
        return text, metadata
    
    def _clean_text(self, text: str) -> str:
        """
        Clean extracted text by removing extra whitespace, fixing line breaks, etc.
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        if not text:
            return ""
            
        # Replace multiple spaces with a single space
        text = re.sub(r'\s+', ' ', text)
        
        # Fix common OCR/extraction issues
        text = re.sub(r'([a-z])- ([a-z])', r'\1\2', text)  # Fix hyphenated words
        
        # Fix paragraph boundaries
        text = re.sub(r'\. ([A-Z])', r'.\n\1', text)  # Add newline after periods followed by capital letters
        
        # Remove redundant newlines but preserve paragraph structure
        text = re.sub(r'\n+', '\n', text)
        
        return text.strip()
    
    def _identify_sections(self, text: str) -> Dict[str, str]:
        """
        Identify common sections in research papers.
        
        Args:
            text: Cleaned full text
            
        Returns:
            Dictionary of section name -> content
        """
        # Common section names in research papers
        section_patterns = [
            (r'abstract', 'abstract'),
            (r'introduction', 'introduction'),
            (r'literature review|related work', 'literature_review'),
            (r'methodology|methods|experimental setup', 'methodology'),
            (r'results', 'results'),
            (r'discussion', 'discussion'),
            (r'conclusion', 'conclusion'),
            (r'references|bibliography', 'references')
        ]
        
        sections = {}
        
        # Extract abstract separately using common patterns
        abstract_match = re.search(r'(?i)abstract[:\s]+([\s\S]+?)(?:\n\s*\n|\n\s*[0-9]+\.|\n\s*keywords|\n\s*introduction)', text)
        if abstract_match:
            sections['abstract'] = abstract_match.group(1).strip()
        
        # Look for section headers and extract content
        for pattern, section_key in section_patterns:
            # Skip abstract if already extracted
            if section_key == 'abstract' and 'abstract' in sections:
                continue
                
            regex = re.compile(r'(?i)(?:^|\n)\s*((?:' + pattern + r')[.\s:]+)([\s\S]+?)(?=(?:^|\n)\s*(?:[A-Z][a-z]+\s+)(?:[.\s:]+)|\Z)', re.MULTILINE)
            match = regex.search(text)
            
            if match:
                sections[section_key] = match.group(2).strip()
        
        # Extract references
        if 'references' in sections:
            references_text = sections['references']
            references = []
            
            # Try to split references into individual entries
            ref_candidates = re.split(r'\n(?:\[\d+\]|\d+\.|\[.+?\])', references_text)
            if len(ref_candidates) > 1:
                # First item might be the header
                references = [ref.strip() for ref in ref_candidates[1:] if ref.strip()]
            else:
                # Try other common formats
                ref_candidates = re.split(r'\n(?:[A-Z][a-z]+,\s+[A-Z]\.)', references_text)
                if len(ref_candidates) > 1:
                    references = ["Author, " + ref.strip() for ref in ref_candidates[1:] if ref.strip()]
            
            if references:
                sections['references_list'] = references
        
        return sections

    def extract_abstract(self, text: str) -> Optional[str]:
        """
        Extract abstract from full text.
        
        Args:
            text: Full document text
            
        Returns:
            Abstract text or None if not found
        """
        # First check if we've already identified sections
        sections = self._identify_sections(text)
        if 'abstract' in sections:
            return sections['abstract']
            
        # Otherwise try to extract it directly
        abstract_patterns = [
            r'(?i)abstract[:\s]+([^#]+?)(?:\n\s*\n|$|\n\s*[0-9]+\.|\n\s*keywords|\n\s*introduction)',
            r'(?i)abstract[:\s]+([\s\S]+?)(?:\n\s*\n|$|\n\s*[0-9]+\.|\n\s*keywords|\n\s*introduction)'
        ]
        
        for pattern in abstract_patterns:
            matches = re.search(pattern, text)
            if matches:
                return matches.group(1).strip()
                
        return None